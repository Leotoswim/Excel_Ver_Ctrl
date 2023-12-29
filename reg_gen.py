##!/projects2/CAD/mcu/script/python/3.6.13/bin/python
# -*- coding: UTF-8 -*- {{{

"""
Generate register related files.
"""

__author__ = [
    '"xiaoqiang" <xiaoqiang@signoff.cn>',
]
__version__='0.10'

#=======================================================================
#
# Created by         : GREAT_USER
# Filename           : reg_gen.py
# Created On         : 2021-6-20 15:00:00
# Description        : Regsiter related files generator
#
#=======================================================================}}}

#import {{{
import os
import sys
import argparse
import json
import re
import xlrd
import time
import copy
from docx import Document
from docx.shared import Cm, RGBColor,Pt
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX,WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from jinja2 import Template
from jinja2 import PackageLoader,Environment,FileSystemLoader

#if os.getenv("DIG_CAD_PATH"):
#    DIG_CAD_PATH = "/projects2/CAD/mcu"
#else:
#    DIG_CAD_PATH = os.getenv("DIG_CAD_PATH")
#sys.path.append("%s/script/python/common_lib/zflow_common"%(DIG_CAD_PATH))

from chip_comlib import CCHIPObject, CCHIPLog, chipTimeStamp, chipUser, chipDate, chipPPformat, chipExitAll, chipMkDir
# }}}

def myGetOpt(): #{{{
    parser = argparse.ArgumentParser()
    parser.add_argument('-i', '--in_file'     , action = 'store',      help = 'Input register excel file'   , default='')
    parser.add_argument('-doc', '--doc'     , action = 'store_true',      help = 'generate document or not '   , default=False)
    parser.add_argument('-o', '--out_dir'     , action = 'store',      help = 'Output directory', default='./out')
    parser.add_argument('-d', '--debug_en' , action = 'store_true', help = 'Debug enable'      , default = False) #must have !!! PLEASE KEEP!
    parser.add_argument('-t', '--trim_cfg' , action = 'store', help = 'trim cfg file'      , default = "") 

    args = parser.parse_args()

    if args.debug_en:
        print(args)

    return args
#}}}


def procTemplate(file_name,template_name,render_dict):#{{{
    template_root=os.getenv('MCU_BU_PythonScript_SharePoint')+"\\Code\\RegGen\\template"
    env = Environment(loader=FileSystemLoader(template_root),keep_trailing_newline=False,lstrip_blocks=True,trim_blocks=False)
    template = env.get_template(template_name)
    output_lines = template.render(render_dict)
    #output file
    fh = open(file_name, "w")
    fh.write(output_lines)
    fh.close()
#}}}


class RegGen(): #{{{
    """
        Main Process Class of RegGen
    """
    def __init__(self, opts,log): #{{{
        self.opts      = opts
        self.in_file   = opts.in_file
        self.out_dir   = opts.out_dir
        self.log       = log
        self.dbg       = opts.debug_en
        self.doc       = opts.doc
        self.project_name = ""
        self.excel_dict = {}
        self.time_tuple = time.localtime(time.time())
    #}}}

    def getExcelInfo(self):#{{{
        self.log.info("getExcelInfo is running !")
        if not os.path.isfile(self.in_file):
            self.log.fatal("file '%s' not existing, please check!!!" %(self.in_file))
            chipExitAll()
        book = xlrd.open_workbook(self.in_file)
        sheets = book.sheet_names()
        for sheet in sheets:
            if sheet == "Revision_History": continue
            table = book.sheet_by_name(sheet)
            row_num = table.nrows
            col_num = table.ncols
            self.excel_dict[sheet] = {"title_dict":{},"title_idx_dict":{},"register_dict":{},"register_name_list":[]}
            self.excel_dict[sheet]["title_dict"]["BASE Address"] = 0
            title_info_scan_type = 0 #0:search project info, 1:search col title idx,  2: search register info
            check_field_bit_info_list = []
            for row_idx in range(0,row_num):
                if(table.cell_value(row_idx,0) == "#REGISTER_DEFINE#"):
                    title_info_scan_type = 1

                #search project info
                if(table.cell_value(row_idx,0) !="" and title_info_scan_type == 0):
                    self.excel_dict[sheet]["title_dict"][table.cell_value(row_idx,0).strip()] = str(table.cell_value(row_idx,1)).strip()
                    if table.cell_value(row_idx,0).strip() == "Base Address":
                        self.excel_dict[sheet]["title_dict"]["Base Address"] = int(re.sub("0x","",self.excel_dict[sheet]["title_dict"]["Base Address"].strip().replace("_",""),re.I),16)
                        if self.excel_dict[sheet]["title_dict"]["Base Address"] > 0x10000:
                            self.log.warn(("Base Address (%d,1) is greater than 0x10000, plase do not add this module's base address in this cell, this is only the offset between different rgf in your IP.")%(row_idx+1))


                #search title col idx
                if(table.cell_value(row_idx,0) == "Offset" and title_info_scan_type == 1):
                    for col_idx in range (0,col_num):
                        if(table.cell_value(row_idx,col_idx) != ""):
                            self.excel_dict[sheet]["title_idx_dict"][table.cell_value(row_idx,col_idx).strip()] = col_idx
                    title_info_scan_type = 2
                    continue

                #search register info
                if(title_info_scan_type == 2):
                    #print(row_idx)
                    #print(table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Offset"])) 
                    #print(table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Reg Name"]))
                    #print(table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Width"]))
                    #print(table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field Name"]))
                    #print(table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field"]))

                    if ((table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Offset"]) != "" and table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Width"]) != "") and
                    (table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Reg Name"]) != "" or table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Long Reg Name"]) != "") and 
                    (table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field Name"]) != "" or table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Long Field Name"]) != "")and 
                    re.match("\[\d+\s*(:\s*\d+)?\]",table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field"]))):

                        if len(check_field_bit_info_list) != 0 and len(check_field_bit_info_list) != self.excel_dict[sheet]["register_dict"][current_reg_name]["Width"]:
                            #print(check_field_bit_info_list)
                            self.log.fatal("Wrong format in (%s,%s), bit gap exist in register %s, please check!!!" %(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field"],current_reg_name))
                            chipExitAll()

                        if table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Reg Name"]).strip() != "": 
                            current_reg_name  = table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Reg Name"]).strip().replace(" ","_").replace("-","_")
                        else:
                            current_reg_name  = table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Long Reg Name"]).strip().replace(" ","_").replace("-","_")
                        self.excel_dict[sheet]["register_name_list"].append(current_reg_name)
                        if "Long Reg Name" in self.excel_dict[sheet]["title_idx_dict"].keys():
                            long_reg_name = table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Long Reg Name"]).strip().replace(" ","_").replace("-","_")
                        else:
                            long_reg_name = current_reg_name

                        if "Header Description" in self.excel_dict[sheet]["title_idx_dict"].keys():
                            header_description = table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Header Description"]).strip().replace("-","_")
                        else:
                            header_description = ""

                        if "Customer Visibility" in self.excel_dict[sheet]["title_idx_dict"].keys():
                            customer_visibility = table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Customer Visibility"]).strip().replace("\n","")
                        else:
                            customer_visibility = ""

                        self.excel_dict[sheet]["register_dict"][current_reg_name] = {
                                #"Offset":int(table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Offset"]).strip().replace("0x","").replace("0X",""),16)+self.excel_dict[sheet]["title_dict"]["Base Address"],
                                "Offset":int(table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Offset"]).strip().replace("0x","").replace("0X",""),16),
                                "Width":int(str(table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Width"])).strip().replace(".0","")),
                                "Reg_Description":table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Reg Description"]).strip().replace("\n",""),
                                "Long_Reg_Name": long_reg_name,
                                 "Header_Description": header_description,
                                "Customer_Visibility": customer_visibility,
                                "field_info_list":[]
                                }

                        if self.excel_dict[sheet]["register_dict"][current_reg_name]["Long_Reg_Name"] == "":
                            self.excel_dict[sheet]["register_dict"][current_reg_name]["Long_Reg_Name"] = current_reg_name

                        check_field_bit_info_list = []
                    elif ((table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Offset"]) == "" and table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Reg Name"]) == "" and table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Width"]) == "") and table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field Name"]) != "" and re.match("\[\d+\s*(:\s*\d+)?\]",table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field"]))):
                        pass
                    else:
                        #self.log.fatal("Wrong format in row %s !!!" %(row_idx))
                        #this is commnet line
                        #print("contine %d"%(row_idx))
                        continue


                    size = 1
                    lsb = 0
                    #check field bit info
                    m = re.match(r"\[\s*(\d+)\s*\]",table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field"]).strip())
                    if m:
                        if m.group(1) in check_field_bit_info_list:
                            self.log.fatal("Wrong format in (%s,%s), bit overlap, please check!!!" %(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field"]))
                            chipExitAll()
                        else:
                            check_field_bit_info_list.append(m.group(1))
                            size = 1
                            lsb = m.group(1)
                    else:
                        m = re.match(r"\[\s*(\d+)\s*\:\s*(\d+)\s*\]",table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field"]).strip())
                        if m:
                            msb = int(m.group(1))
                            lsb = int(m.group(2))
                            if msb < lsb:
                                self.log.fatal("Wrong format in (%s,%s), lsb > msb, please check!!!" %(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field"]))
                                chipExitAll()
                            else:
                                size = msb - lsb + 1
                                for i in range(lsb,msb+1):
                                    if str(i) in check_field_bit_info_list:
                                        self.log.fatal("Wrong format in (%s,%s), bit overlap, please check!!!" %(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field"]))
                                        chipExitAll()
                                    else:
                                        check_field_bit_info_list.append(str(i))
                        else:
                            self.log.fatal("Wrong format in (%s,%s), field info, please check!!!" %(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field"]))
                            chipExitAll()

                    #check the last register bit gap exist or not
                    if row_idx == row_num -1 and len(check_field_bit_info_list) != 0 and len(check_field_bit_info_list) !=  self.excel_dict[sheet]["register_dict"][current_reg_name]["Width"]:
                        self.log.fatal("Wrong format in (%s,%s), bit gap exist in register %s, please check!!!" %(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field"],current_reg_name))
                        chipExitAll()
                    #print(row_idx)
                    #print(table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field Name"]).strip())
                    if(re.match("^\(?reserved\)?$",table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field Name"]).strip(),re.I)):
                        continue
                    

                    if re.match("0|0.0",str(table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Reset Value"])).strip()):
                        reset_value = 0
                    elif re.match("\d*'b[0-1]+",str(table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Reset Value"])).replace(".0","").strip()):
                        reset_value = int(re.sub(".*'b","",str(table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Reset Value"])).replace(".0","").strip()), 2)
                    elif re.match("\d*'h[0-9a-fA-F]+",str(table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Reset Value"])).replace(".0","").strip()):
                        reset_value = int(re.sub(".*'h","",str(table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Reset Value"])).replace(".0","").strip()),16)
                    else:
                        self.log.fatal("Wrong format in (%s,%s), reset value format must be like 3'h4 or 2'b11, please check!!!" %(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Reset Value"]))
                        chipExitAll()
                        reset_value =0
                    reset_value_b = bin(reset_value).replace("0b","")
                    reset_value_l = []
                    for i in range(size - len(reset_value_b)):
                        reset_value_l.append(0)
                    for i in reset_value_b:
                        reset_value_l.append(int(i))
                    reset_value_l.reverse()

                    if "Long Field Name" in self.excel_dict[sheet]["title_idx_dict"].keys():
                        long_field_name = table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field Name"]).strip().replace(" ","_").replace("-","_")
                    else:
                        long_field_name = ""

                    if "Header Description" in self.excel_dict[sheet]["title_idx_dict"].keys():
                        #header_description = table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Header Description"]).strip()
                        header_description = table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Header Description"]).strip().translate(str.maketrans("‘’，：–\"“”…；×（）","'',:_''' ;x()")),
                        header_description = header_description[0].replace("-","_")
                    else:
                        header_description = ""

                    if "Customer Visibility" in self.excel_dict[sheet]["title_idx_dict"].keys():
                        customer_visibility = table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Customer Visibility"]).strip().replace("\n","")
                    else:
                        customer_visibility = ""
                        
                    if (table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field Description"]).strip() == ""):
                        field_description = table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field Name"]).strip() 
                    else:
                        field_description = table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field Description"]).strip().translate(str.maketrans("。‘’，：–\"“”…；×（）℃",".'',:_''' :x() "))

                    field_dict = {
                            "Field Name":table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field Name"]).strip().replace(" ","_").replace("-","_"),
                            "Long Field Name":long_field_name,
                            "Field":table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field"]).strip(),
                            "Access Right":table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Access Right"]).strip(),
                            "Reset Value":reset_value,
                            "reset_value_l":reset_value_l,
                            "Set/Clear":table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Set/Clear"]).strip(),
                            #"Field_Description":table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field Description"]).strip().replace(":—",": ").replace(";—","; ").translate(str.maketrans("‘’，：–\"“”…；×（）℃","'',:_''' :x() ")),
                            #"Field_Description":table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Field Description"]).strip(),
                            "Field_Description":field_description,
                            "Header_Description":header_description,
                            #"Header_Description":table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Header Description"]).strip().translate(str.maketrans("‘’，：–\"“”…；×（）","'',:_''' ;x()")),
                            "Customer_Visibility":customer_visibility,
                            "size":size,
                            "lsb":int(lsb)
                            }
                    if table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Misc"]).strip() == "WR_HIT_OUT" or table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Misc"]).strip() == "WR_HITS_OUT":
                        self.excel_dict[sheet]["register_dict"][current_reg_name]["WR_HIT_OUT"] = True
                    if table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Misc"]).strip() == "WR_HIT_OUT_DLY" or table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Misc"]).strip() == "WR_HITS_OUT_DLY":
                        self.excel_dict[sheet]["register_dict"][current_reg_name]["WR_HIT_OUT_DLY"] = True
                    if table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Misc"]).strip() == "RD_HIT_OUT" or table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Misc"]).strip() == "RD_HITS_OUT":
                        self.excel_dict[sheet]["register_dict"][current_reg_name]["RD_HIT_OUT"] = True
                    if table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Misc"]).strip() == "RD_HIT_OUT_DLY" or table.cell_value(row_idx,self.excel_dict[sheet]["title_idx_dict"]["Misc"]).strip() == "RD_HITS_OUT_DLY":
                        self.excel_dict[sheet]["register_dict"][current_reg_name]["RD_HIT_OUT_DLY"] = True
                    if field_dict["Access Right"] == "RW_E" or field_dict["Access Right"] == "WR_E" or field_dict["Access Right"] == "WC_E" or field_dict["Access Right"] == "RC_W1E" or field_dict["Access Right"] == "RC_W0E" or field_dict["Access Right"] == "RC_RE" or  field_dict["Access Right"] == "RS_E": 
                        self.excel_dict[sheet]["register_dict"][current_reg_name]["WR_HIT_OUT"] = True
                    if field_dict["Long Field Name"] == "" and field_dict["Field Name"] != "":
                        field_dict["Long Field Name"] = field_dict["Field Name"]
                    if field_dict["Field Name"] == "" and field_dict["Long Field Name"] != "":
                        field_dict["Field Name"] = field_dict["Long Field Name"]
                    self.excel_dict[sheet]["register_dict"][current_reg_name]["field_info_list"].append(field_dict)


        self.log.debug(str(self.excel_dict))
        self.chkExcelFormat()
        #print(str(self.excel_dict))
    #}}}

    def chkExcelFormat(self): #{{{
        for sheet in sorted(self.excel_dict.keys()):
            address_byte_list = []
            current_width = -1
            for register in self.excel_dict[sheet]["register_name_list"]:
                width = self.excel_dict[sheet]["register_dict"][register]["Width"]
                offset = self.excel_dict[sheet]["register_dict"][register]["Offset"]
                byte_num = int(width/8)

                #check offset align with width
                if("Offset Format" in self.excel_dict[sheet]["title_dict"] and self.excel_dict[sheet]["title_dict"]["Offset Format"] != "Series"):
                    if offset%byte_num != 0:
                        self.log.fatal("Offset for register %s not aligned with register width, please check!!!"%(register))
                        chipExitAll()

                #check address overlap
                if("Offset Format" in self.excel_dict[sheet]["title_dict"] and self.excel_dict[sheet]["title_dict"]["Offset Format"] != "Series"):
                    for i in range(0,byte_num):
                        if (offset + i) in address_byte_list:
                            self.log.fatal("Offset for register %s overlapped with other register, please check!!!"%(register))
                            chipExitAll()
                        else:
                            address_byte_list.append(offset+i)

                #check register width all the same
                if current_width == -1:
                    current_width = width
                else:
                    if not current_width == width:
                        self.log.fatal("register width not all the same, current register is %s, width is %s, please check!!!"%(register,width))
                        chipExitAll()
            self.excel_dict[sheet]["reg_width"] = current_width



                #for field_dict in self.excel_dict[sheet]["register_dict"][register]["field_info_list"]:
                #    if field_dict['Long Field Name'] != "":
                #        output_line += " "*4 + """rand uvm_reg_field %s;\n"""%(field_dict["Long Field Name"])

    #}}}

    def genRegModel(self):#{{{
        self.log.debug("genRegModel is running !")
        for sheet in sorted(self.excel_dict.keys()):
            chipMkDir("%s/%s"%(self.out_dir,sheet))
            module_name  = self.excel_dict[sheet]["title_dict"]["Module Name"]
            output_line = ""
            for register in self.excel_dict[sheet]["register_name_list"]:
                if self.excel_dict[sheet]["register_dict"][register]["Long_Reg_Name"] != "":
                    register_name = self.excel_dict[sheet]["register_dict"][register]["Long_Reg_Name"].lower()
                else:
                    register_name = register.lower() 
                output_line +="""class %s_%s extends easy_uvm_reg;\n"""%(module_name,register_name)
                for field_dict in self.excel_dict[sheet]["register_dict"][register]["field_info_list"]:
                    if field_dict['Long Field Name'] != "":
                        output_line += " "*4 + """rand uvm_reg_field %s;\n"""%(field_dict["Long Field Name"])
                    else:
                        output_line += " "*4 + """rand uvm_reg_field %s;\n"""%(field_dict["Field Name"])

                output_line += "\n" + " "*4 + """`uvm_object_utils(%s_%s)\n
    function new(string name = "%s_%s");
        super.new(name, %d, build_coverage(UVM_NO_COVERAGE));
    endfunction: new\n
    virtual function void build();\n"""%(module_name,register_name,module_name,register_name,self.excel_dict[sheet]["register_dict"][register]["Width"])

                for field_dict in self.excel_dict[sheet]["register_dict"][register]["field_info_list"]:
                    if field_dict['Long Field Name'] != "":
                        output_line += " "*8 + """this.%s = uvm_reg_field::type_id::create("%s",,get_full_name());\n"""%(field_dict["Long Field Name"].ljust(20),field_dict["Long Field Name"])
                    else:
                        output_line += " "*8 + """this.%s = uvm_reg_field::type_id::create("%s",,get_full_name());\n"""%(field_dict["Field Name"].ljust(20),field_dict["Field Name"])

                for field_dict in self.excel_dict[sheet]["register_dict"][register]["field_info_list"]:
                    if(field_dict["Access Right"] =="WR_E" or field_dict["Access Right"] =="WC_E" or  field_dict["Access Right"] =="WC_G" or
                        field_dict["Access Right"] =="WR_G" or field_dict["Access Right"] =="RW_G" or  field_dict["Access Right"] =="RC_W1E" or
                        field_dict["Access Right"] =="RC_W0E" or field_dict["Access Right"] =="RC_RE" or  field_dict["Access Right"] =="RC_W0E" or
                        field_dict["Access Right"] =="RS_E" or field_dict["Access Right"] =="RW_E"):
                        access_right = "RW"
                    else:
                        access_right = field_dict["Access Right"]
                    if field_dict['Long Field Name'] != "":
                        output_line += " "*8 + """this.%s.configure(this,%d,%d,"%s",0,'h%x,1,0,1);\n"""%(field_dict["Long Field Name"],field_dict["size"],field_dict["lsb"],access_right,field_dict["Reset Value"])
                    else:
                        output_line += " "*8 + """this.%s.configure(this,%d,%d,"%s",0,'h%x,1,0,1);\n"""%(field_dict["Field Name"],field_dict["size"],field_dict["lsb"],access_right,field_dict["Reset Value"])

                output_line += " "*4 + """endfunction\nendclass\n\n\n"""

            output_line +="""class %s_reg_model extends uvm_reg_block;\n"""%(module_name)
            for register in self.excel_dict[sheet]["register_name_list"]:
                if self.excel_dict[sheet]["register_dict"][register]["Long_Reg_Name"] != "":
                    register_name = self.excel_dict[sheet]["register_dict"][register]["Long_Reg_Name"].lower()
                else:
                    register_name = register.lower() 
                output_line += " "*4 + """rand %s_%s %s;\n"""%(module_name,register_name.ljust(20),register_name)

            output_line +="\n" + " "*4 + """`uvm_object_utils(%s_reg_model)\n
    uvm_reg_map         reg_map  ;\n
    function new(string name = "%s_reg_model");
      super.new(name, build_coverage(UVM_NO_COVERAGE));
    endfunction \n
    virtual function void build();
        this.reg_map = create_map("reg_map", 0, 4, UVM_LITTLE_ENDIAN, 0);\n\n"""%(module_name,module_name)

            for register in self.excel_dict[sheet]["register_name_list"]:
                output_line += " "*8 + """this.%s = %s_%s::type_id::create("%s",,get_full_name());\n"""%(register_name,module_name,register_name,register_name)
                output_line += " "*8 + """this.%s.configure(this,null,"");\n"""%(register_name)
                output_line += " "*8 + """this.%s.build();\n\n"""%(register_name)

            for register in self.excel_dict[sheet]["register_name_list"]:
                output_line += " "*8 + """this.reg_map.add_reg(this.%s,`UVM_REG_ADDR_WIDTH'h%04x,"RW",0);\n"""%(register_name,self.excel_dict[sheet]["register_dict"][register]["Offset"])

            output_line += """
        lock_model();
    endfunction
	
    function void set_frontdoor(uvm_reg_frontdoor ftdr);
        uvm_reg regs[$];
        this.get_registers(regs);
        foreach(regs[i])begin
            regs[i].set_frontdoor(ftdr);
        end
    endfunction: set_frontdoor\n\nendclass\n"""

            fh = open("%s/%s/%s_reg_model.sv"%(self.out_dir,sheet,module_name),"w")
            fh.write(output_line)
            fh.close()
            #os.system("rm -f %s/%s/%s_rgf_uvm_reg_model.sv"%(self.out_dir,sheet,module_name))
    #}}}

    def genRegModelLite(self):#{{{
        self.log.debug("genRegModelSv is running !")
        for sheet in sorted(self.excel_dict.keys()):
            chipMkDir("%s/%s"%(self.out_dir,sheet))
            module_name  = self.excel_dict[sheet]["title_dict"]["Module Name"]
            output_line = ""
            for register in self.excel_dict[sheet]["register_name_list"]:
                if self.excel_dict[sheet]["register_dict"][register]["Long_Reg_Name"] != "":
                    register_name = self.excel_dict[sheet]["register_dict"][register]["Long_Reg_Name"].lower()
                else:
                    register_name = register.lower() 
                output_line += "="*30+register_name+"="*30+"\n"
                output_line +="""reg_dict["%s_%s"]["__address__"]["addr"] = 32'h%x;\n"""%(module_name,register_name,self.excel_dict[sheet]["register_dict"][register]["Offset"])

                for field_dict in self.excel_dict[sheet]["register_dict"][register]["field_info_list"]:
                    if field_dict['Long Field Name'] != "":
                        field_name = field_dict["Long Field Name"]
                    else:
                        field_name = field_dict["Field Name"]
                    output_line +="""reg_dict["%s_%s"]["%s"]["lsb"]   = %d;\n"""%(module_name,register_name,field_name,field_dict["lsb"])
                    output_line +="""reg_dict["%s_%s"]["%s"]["size"]  = %d;\n"""%(module_name,register_name,field_name,field_dict["size"])
                    output_line +="""reg_dict["%s_%s"]["%s"]["reset"] = 'h%x;\n"""%(module_name,register_name,field_name,field_dict["Reset Value"])
                    output_line +="""reg_dict["%s_%s"]["%s"]["type"]  = '%s';\n"""%(module_name,register_name,field_name,field_dict["Access Right"])
            fh = open("%s/%s/%s_reg_model_lite.sv"%(self.out_dir,sheet,module_name),"w")
            fh.write(output_line)
            fh.close()
            #os.system("rm -f %s/%s/%s_rgf_uvm_reg_model.sv"%(self.out_dir,sheet,module_name))
    #}}}

    def genRegDef(self):#{{{
        self.log.debug("genRegDef is running !")
        for sheet in sorted(self.excel_dict.keys()):
            chipMkDir("%s/%s"%(self.out_dir,sheet))
            module_name  = self.excel_dict[sheet]["title_dict"]["Module Name"]
            output_line = """`ifndef %s_BASE_ADDR \n    `define %s_BASE_ADDR 'h0\n`endif\n\n"""%(self.excel_dict[sheet]["title_dict"]["Module Name"].upper(),self.excel_dict[sheet]["title_dict"]["Module Name"].upper())
            for register in self.excel_dict[sheet]["register_name_list"]:
                if self.excel_dict[sheet]["register_dict"][register]["Long_Reg_Name"] != "":
                    register_name = self.excel_dict[sheet]["register_dict"][register]["Long_Reg_Name"].lower()
                else:
                    register_name = register.lower() 
                output_line +="""`define %s_%s `%s_BASE_ADDR + 'h%04x \n"""%(self.excel_dict[sheet]["title_dict"]["Module Name"].upper(),register_name.upper(),self.excel_dict[sheet]["title_dict"]["Module Name"].upper(),self.excel_dict[sheet]["register_dict"][register]["Offset"])
            fh = open("%s/%s/%s_reg_def.v"%(self.out_dir,sheet,module_name),"w")
            fh.write(output_line)
            fh.close()
            #os.system("rm -f %s/%s/%s_rgf_def.v"%(self.out_dir,sheet,module_name))
    #}}}

    def genRegDef2(self):#{{{
        self.log.debug("genRegDef2 is running !")
        for sheet in sorted(self.excel_dict.keys()):
            chipMkDir("%s/%s"%(self.out_dir,sheet))
            module_name  = self.excel_dict[sheet]["title_dict"]["Project Name"]
            output_line = "/*------------- %s -----------*/\n"%module_name.upper()

            output_line += """\n"""
            if not ("Offset Format" in self.excel_dict[sheet]["title_dict"] and self.excel_dict[sheet]["title_dict"]["Offset Format"] == "Series"):
                output_line += "typedef struct\n"
                output_line += "{\n"
                last_offset = -1
                for register in self.excel_dict[sheet]["register_name_list"]:
                    offset = self.excel_dict[sheet]["register_dict"][register]["Offset"]
                    width  = self.excel_dict[sheet]["register_dict"][register]["Width"]
                    if width != 32:
                        break
                    if(offset%4!=0):
                        self.log.warn("offset is not aligned with register width! will not generate %s_reg_def.h",module_name)
                        break
                    while((last_offset != -1) and (last_offset+4 != offset)):
                        last_offset += 4
                        output_line +=("""  __IO    uint32_t """.ljust(20)) + "RESERVED_%s;\n"%(hex(last_offset))
                    output_line +=(("""  __IO    uint32_t %s;"""%(register.upper())).ljust(40)) + "//0x%04x\n"%(self.excel_dict[sheet]["register_dict"][register]["Offset"])
                    last_offset = offset
                output_line += "\n"
                output_line += "} %s_TypeDef;\n"%self.excel_dict[sheet]["title_dict"]["Project Name"].upper()
            else:
                for register in self.excel_dict[sheet]["register_name_list"]:
                    output_line +=(("""#define %s """%(register.upper())).ljust(40)) + "0x%04x\n"%(self.excel_dict[sheet]["register_dict"][register]["Offset"])
            output_line += "\n"
            
            for register in self.excel_dict[sheet]["register_name_list"]:
                for field in self.excel_dict[sheet]["register_dict"][register]["field_info_list"]:
                    output_line +=("""#define %s_%s_%s_Pos"""%(self.excel_dict[sheet]["title_dict"]["Project Name"].upper(),register.upper(),field["Field Name"].upper())).ljust(80) + "%d\n"%(field["lsb"])
                    output_line +=("""#define %s_%s_%s_Msk"""%(self.excel_dict[sheet]["title_dict"]["Project Name"].upper(),register.upper(),field["Field Name"].upper())).ljust(80) + "(0x%xUL <<  %s_%s_%s_Pos)\n"%((2**(int(field["size"]))-1),self.excel_dict[sheet]["title_dict"]["Project Name"].upper(),register.upper(),field["Field Name"].upper())

                    if (field["Header_Description"] != ""):
                        postfixs = field["Header_Description"].split(";")
                        for postfix in postfixs:
                            m = re.match("\s*(\S+)\s*:\s*(\S+)\s*",postfix)
                            if m:
                                if re.match("[0x|0d]\S+",m.group(1),re.I):
                                    output_line +=("""#define %s_%s_%s_%s"""%(self.excel_dict[sheet]["title_dict"]["Project Name"].upper(),register.upper(),field["Field Name"].upper(),m.group(2).replace("/","_").upper())).ljust(80) + m.group(1) + "\n"
                                else:
                                    output_line +=("""#define %s_%s_%s_%s"""%(self.excel_dict[sheet]["title_dict"]["Project Name"].upper(),register.upper(),field["Field Name"].upper(),m.group(2).replace("/","_").upper())).ljust(80) +   "0x" + m.group(1) + "\n"
                            #else:
                            #    print("[ERROR] format for Header Description %s"%field["Header_Description"])

                    output_line +="\n"
            fh = open("%s/%s/%s_reg_def.h"%(self.out_dir,sheet,module_name),"w")
            fh.write(output_line)
            fh.close()
    #}}}
    
    def genRegSvd(self):#{{{
        self.log.debug("genRegSvd is running !")
        for sheet in sorted(self.excel_dict.keys()):
            chipMkDir("%s/%s"%(self.out_dir,sheet))
            module_name  = self.excel_dict[sheet]["title_dict"]["Project Name"]
            #print(self.excel_dict)
            output_line = ""
            for register in self.excel_dict[sheet]["register_name_list"]:

                #if all fields are not customer visibile, then ignore this register
                register_customer_visibility = False
                for field in self.excel_dict[sheet]["register_dict"][register]["field_info_list"]:
                    if field["Customer_Visibility"] != "N":
                        register_customer_visibility = True 
                        break
                if register_customer_visibility == False:
                    continue

                output_line += "            <register>\n"
                output_line += "               <name>%s</name>\n"%register.upper()
                #output_line += "               <displayName>%s</displayName>\n"%register.upper()
                output_line += "               <description>%s</description>\n"%self.excel_dict[sheet]["register_dict"][register]["Long_Reg_Name"]
                output_line += "               <addressOffset>0x%04x</addressOffset>\n"%self.excel_dict[sheet]["register_dict"][register]["Offset"]
                output_line += "               <size>%d</size>\n"%self.excel_dict[sheet]["register_dict"][register]["Width"]
                reset_value = 0
                for field in self.excel_dict[sheet]["register_dict"][register]["field_info_list"]:
                    reset_value = reset_value | field["Reset Value"]<<int(re.sub("\[(\d+:)?","",field["Field"].replace("]","")))
                output_line += "               <resetValue>0x%x</resetValue>\n"%(reset_value)   

                output_line += "               <fields>\n"
                for field in self.excel_dict[sheet]["register_dict"][register]["field_info_list"]:
                    #if this field is not customer visibile, then ignore this field
                    if field["Customer_Visibility"] == "N":
                        continue
                    output_line +="                  <field>\n"
                    output_line +="                     <name>%s</name>\n"%field["Field Name"].upper()
                    if re.match("reserved",field["Field_Description"],re.I):
                        output_line +="                     <description>%s</description>\n"%field["Field Name"]
                    else:
                        output_line +="                     <description>%s</description>\n"%field["Field_Description"]
                    if field["size"] != 1:
                        output_line +="                     <bitRange>%s</bitRange>\n"%(field["Field"].replace(" ",""))
                    else:
                        output_line +="                     <bitRange>%s</bitRange>\n"%((field["Field"]+field["Field"]).replace("][",":"))

                    if re.match("RW|WR|RW_E|WR_G|RW_G|RC_W1E|RC_W0E|RS_E|WR_E",field["Access Right"],re.I):
                        output_line +="                     <access>%s</access>\n"%("read-write")
                    elif re.match("RO|RC_RE",field["Access Right"],re.I):  
                        output_line +="                     <access>%s</access>\n"%("read-only")
                    elif re.match("WO|WC|WC_E|WC_G",field["Access Right"],re.I):
                        output_line +="                     <access>%s</access>\n"%("write-only")
                    else:
                        output_line +="                     <access>%s</access>\n"%field["Access Right"]

                    if (field["Header_Description"] != ""):
                        output_line +="                     <enumeratedValues>\n"
                        field["Header_Description"] = re.sub("\n",";", field["Header_Description"])
                        postfixs = field["Header_Description"].split(";")
                        postfixs = field["Header_Description"].split(";")
                        for postfix in postfixs:
                            postfix = postfix.strip()
                            if postfix == "": continue
                            output_line +="                         <enumeratedValue>\n"
                            m = re.match("(.*):(.*)",postfix)
                            if m:
                                value = re.sub("/","_",m.group(1).strip())
                                name = re.sub("/","_",m.group(2).strip())
                                name = re.sub(" ","_",name)
                                if(re.match("reserved",name,re.I)):
                                    output_line +="                             <name>reserved_not_used</name>\n"
                                    output_line +="                             <description>reserved_not_used</description>\n"
                                else:
                                    output_line +="                             <name>%s</name>\n"%(name.upper())
                                    output_line +="                             <description>%s</description>\n"%(name.upper())
                                if re.match("0x\S+",value,re.I):
                                    output_line +="                             <value>%s</value>\n"%(value)
                                elif re.match("0b\S+",value,re.I):
                                    output_line +="                             <value>0x%x</value>\n"%(int(value.replace("0b",""),2))
                                elif re.match("0d\S+",value,re.I):
                                    output_line +="                             <value>0x%x</value>\n"%(int(value.replace("0d",""),10))
                                elif re.match(".*'h\S+",value,re.I):
                                    output_line +="                             <value>0x%x</value>\n"%(int(re.sub(".*'h","",value)))
                                elif re.match(".*'b\S+",value,re.I):
                                    output_line +="                             <value>0x%x</value>\n"%(int(re.sub(".*'b","",value),2))
                                elif re.match(".*'d\S+",value,re.I):
                                    output_line +="                             <value>0x%x</value>\n"%(int(re.sub(".*'d","",value),10))
                                else:
                                    output_line +="                             <value>%s</value>\n"%("0x" + value)
                            else:
                                print(postfix)
                                print("[ERROR] format for Header Description %s"%field["Header_Description"])
                            output_line +="                         </enumeratedValue>\n"
                        output_line +="                     </enumeratedValues>\n"
                    output_line +="                  </field>\n"

                output_line += "               </fields>\n"
                output_line += "            </register>\n"
            fh = open("%s/%s/%s_reg.svd"%(self.out_dir,sheet,module_name),"w")
            fh.write(output_line)
            fh.close()
    #}}}

    def genRegSvdInternal(self):#{{{
        self.log.debug("genRegSvdFull is running !")
        for sheet in sorted(self.excel_dict.keys()):
            chipMkDir("%s/%s"%(self.out_dir,sheet))
            module_name  = self.excel_dict[sheet]["title_dict"]["Project Name"]
            output_line = ""
            for register in self.excel_dict[sheet]["register_name_list"]:
                output_line += "            <register>\n"
                output_line += "               <name>%s</name>\n"%register.upper()
                #output_line += "               <displayName>%s</displayName>\n"%register.upper()
                output_line += "               <description>%s</description>\n"%self.excel_dict[sheet]["register_dict"][register]["Long_Reg_Name"]
                output_line += "               <addressOffset>0x%04x</addressOffset>\n"%self.excel_dict[sheet]["register_dict"][register]["Offset"]
                output_line += "               <size>%d</size>\n"%self.excel_dict[sheet]["register_dict"][register]["Width"]
                reset_value = 0
                for field in self.excel_dict[sheet]["register_dict"][register]["field_info_list"]:
                    reset_value = reset_value | field["Reset Value"]<<int(re.sub("\[(\d+:)?","",field["Field"].replace("]","")))
                output_line += "               <resetValue>0x%x</resetValue>\n"%(reset_value)   

                output_line += "               <fields>\n"
                for field in self.excel_dict[sheet]["register_dict"][register]["field_info_list"]:
                    #if this field is not customer visibile, then ignore this field
                    output_line +="                  <field>\n"
                    output_line +="                     <name>%s</name>\n"%field["Field Name"].upper()
                    if re.match("reserved",field["Field_Description"],re.I):
                        output_line +="                     <description>%s</description>\n"%field["Field Name"]
                    else:
                        output_line +="                     <description>%s</description>\n"%field["Field_Description"]
                    if field["size"] != 1:
                        output_line +="                     <bitRange>%s</bitRange>\n"%(field["Field"].replace(" ",""))
                    else:
                        output_line +="                     <bitRange>%s</bitRange>\n"%((field["Field"]+field["Field"]).replace("][",":"))

                    if re.match("RW|WR|RW_E|WR_G|RW_G|RC_W1E|RC_W0E|RS_E|WR_E",field["Access Right"],re.I):
                        output_line +="                     <access>%s</access>\n"%("read-write")
                    elif re.match("RO|RC_RE",field["Access Right"],re.I):  
                        output_line +="                     <access>%s</access>\n"%("read-only")
                    elif re.match("WO|WC|WC_E|WC_G",field["Access Right"],re.I):
                        output_line +="                     <access>%s</access>\n"%("write-only")
                    else:
                        output_line +="                     <access>%s</access>\n"%field["Access Right"]

                    if (field["Header_Description"] != ""):
                        output_line +="                     <enumeratedValues>\n"
                        field["Header_Description"] = re.sub("\n",";", field["Header_Description"])
                        postfixs = field["Header_Description"].split(";")
                        postfixs = field["Header_Description"].split(";")
                        for postfix in postfixs:
                            postfix = postfix.strip()
                            if postfix == "": continue
                            output_line +="                         <enumeratedValue>\n"
                            m = re.match("(.*):(.*)",postfix)
                            if m:
                                value = re.sub("/","_",m.group(1).strip())
                                name = re.sub("/","_",m.group(2).strip())
                                name = re.sub(" ","_",name)
                                if(re.match("reserved",name,re.I)):
                                    output_line +="                             <name>reserved_not_used</name>\n"
                                    output_line +="                             <description>reserved_not_used</description>\n"
                                else:
                                    output_line +="                             <name>%s</name>\n"%(name.upper())
                                    output_line +="                             <description>%s</description>\n"%(name.upper())
                                if re.match("0x\S+",value,re.I):
                                    output_line +="                             <value>%s</value>\n"%(value)
                                elif re.match("0b\S+",value,re.I):
                                    output_line +="                             <value>0x%x</value>\n"%(int(value.replace("0b",""),2))
                                elif re.match("0d\S+",value,re.I):
                                    output_line +="                             <value>0x%x</value>\n"%(int(value.replace("0d",""),10))
                                elif re.match(".*'h\S+",value,re.I):
                                    output_line +="                             <value>0x%x</value>\n"%(int(re.sub(".*'h","",value)))
                                elif re.match(".*'b\S+",value,re.I):
                                    output_line +="                             <value>0x%x</value>\n"%(int(re.sub(".*'b","",value),2))
                                elif re.match(".*'d\S+",value,re.I):
                                    output_line +="                             <value>0x%x</value>\n"%(int(re.sub(".*'d","",value),10))
                                else:
                                    output_line +="                             <value>%s</value>\n"%("0x" + value)
                            else:
                                print(postfix)
                                print("[ERROR] format for Header Description %s"%field["Header_Description"])
                            output_line +="                         </enumeratedValue>\n"
                        output_line +="                     </enumeratedValues>\n"
                    output_line +="                  </field>\n"

                output_line += "               </fields>\n"
                output_line += "            </register>\n"
            fh = open("%s/%s/%s_reg_internal.svd"%(self.out_dir,sheet,module_name),"w")
            fh.write(output_line)
            fh.close()
    #}}}

    def genSigBitInfo(self):#{{{
        self.log.debug("genSigBitInfo is running !")
        for sheet in sorted(self.excel_dict.keys()):
            chipMkDir("%s/%s"%(self.out_dir,sheet))
            module_name  = self.excel_dict[sheet]["title_dict"]["Module Name"]
            output_line = ""
            for register in self.excel_dict[sheet]["register_name_list"]:
                if self.excel_dict[sheet]["register_dict"][register]["Long_Reg_Name"] != "":
                    register_name = self.excel_dict[sheet]["register_dict"][register]["Long_Reg_Name"].lower()
                else:
                    register_name = register.lower() 
                for field in self.excel_dict[sheet]["register_dict"][register]["field_info_list"]:
                    if field['Long Field Name'] != "":
                        field_name = field["Long Field Name"]
                    else:
                        field_name = field["Field Name"]
                    output_line +=("""`define SIG_BIT__%s__%s__%s"""%(self.excel_dict[sheet]["title_dict"]["Module Name"].upper(),register_name.upper(),field_name.upper())).ljust(100)+ """%d:%d \n"""%(field["lsb"]+field["size"]-1,field["lsb"])
            fh = open("%s/%s/%s_sig_bit_info.v"%(self.out_dir,sheet,module_name),"w")
            fh.write(output_line)
            fh.close()
    #}}}

    def genDocx(self):#{{{
        self.log.debug("genDocx is running !")
        print("genDocx is running !")
        for sheet in sorted(self.excel_dict.keys()):
            chipMkDir("%s/%s"%(self.out_dir,sheet))
            module_name  = self.excel_dict[sheet]["title_dict"]["Project Name"]

            document = Document()
            section = document.sections[0]
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

            document.add_heading('Register Address Map', level=1)

            self.excel_dict_new_sheet = copy.deepcopy(self.excel_dict[sheet])
            for register_name in self.excel_dict[sheet]["register_name_list"]:
                register_customer_visibility = False
                for field in self.excel_dict[sheet]["register_dict"][register_name]["field_info_list"]:
                    if field["Customer_Visibility"].strip() != "N":
                        register_customer_visibility = True 
                        break
                if register_customer_visibility == False:
                    self.excel_dict_new_sheet["register_name_list"].remove(register_name)
                    self.excel_dict_new_sheet["register_dict"].pop(register_name)
                    continue

            table = document.add_table(rows=1+len(self.excel_dict_new_sheet["register_name_list"]), cols=3,style = 'Table Grid')
            table.autofit = False
            table.allow_autofit = False
            widths =(Inches(1),Inches(1.5),Inches(4.5))
            for row in table.rows:
                for idx,width in enumerate(widths):
                    row.cells[idx].width = width
            hdr_cells = table.rows[0].cells
            hdr_cells[0].paragraphs[0].add_run("Offset").bold = True
            hdr_cells[1].paragraphs[0].add_run("Reg Name").bold = True
            hdr_cells[2].paragraphs[0].add_run("Reg Description").bold = True

            for row_idx in range(0,len(self.excel_dict_new_sheet["register_name_list"])):
                table.rows[row_idx+1].cells[0].text = "0x%04x"%(self.excel_dict_new_sheet["register_dict"][self.excel_dict_new_sheet["register_name_list"][row_idx]]["Offset"])
                table.rows[row_idx+1].cells[1].text = "%s_%s"%(self.excel_dict[sheet]["title_dict"]["Project Name"].upper(),self.excel_dict_new_sheet["register_name_list"][row_idx].upper())
                table.rows[row_idx+1].cells[2].text = self.excel_dict_new_sheet["register_dict"][self.excel_dict_new_sheet["register_name_list"][row_idx]]["Reg_Description"] 

            document.add_heading('Register Field Details', level=1)
            for row_idx in range(0,len(self.excel_dict_new_sheet["register_name_list"])):
                register_name = "%s"%(self.excel_dict_new_sheet["register_name_list"][row_idx])
                document.add_paragraph('', style='List Number').add_run("%s_%s"%(self.excel_dict[sheet]["title_dict"]["Project Name"].upper(),self.excel_dict_new_sheet["register_name_list"][row_idx].upper())).bold = True

                table = document.add_table(rows=9, cols=17,style = 'Table Grid')
                table.alignment = WD_TABLE_ALIGNMENT.CENTER                

                #set cell width; TBD
                table.autofit = False
                table.allow_autofit = False
                widths =(Inches(0.6),Inches(0.4),Inches(0.4),Inches(0.4),Inches(0.4),Inches(0.4),Inches(0.4),Inches(0.4),Inches(0.4),Inches(0.4),Inches(0.4),Inches(0.4),Inches(0.4),Inches(0.4),Inches(0.4),Inches(0.4),Inches(0.4))
                for row in table.rows:
                    for idx,width in enumerate(widths):
                        row.cells[idx].width = width

                table.cell(0,0).text = "0x%04x"%(self.excel_dict_new_sheet["register_dict"][self.excel_dict_new_sheet["register_name_list"][row_idx]]["Offset"])
                table.cell(0,0).merge(table.cell(0,1))
                table.cell(0,0).merge(table.cell(0,2))
                table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for i in range(4,14):
                    table.cell(0,3).merge(table.cell(0,i))
                table.cell(0,3).text = self.excel_dict_new_sheet["register_dict"][register_name]["Reg_Description"]
                table.cell(0,3).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                table.cell(0,14).merge(table.cell(0,15))
                table.cell(0,14).merge(table.cell(0,16))
                table.cell(0,14).text = "%s_%s"%(self.excel_dict[sheet]["title_dict"]["Project Name"].upper(),self.excel_dict_new_sheet["register_name_list"][row_idx].upper())
                table.cell(0,14).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
                table.cell(1,0).text = "Bits"
                table.cell(1,0)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "00FFFF")))
                for i in range(1,17): #31~16
                    table.cell(1,i).text = str(32-i)
                    table.cell(1,i)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "00FFFF")))
                table.cell(2,0).text = "Name"
                table.cell(3,0).text = "Type"
                table.cell(4,0).text = "Reset"
                table.cell(5,0).text = "Bits"
                table.cell(5,0)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "00FFFF")))
                for i in range(1,17): #15~0
                    table.cell(5,i).text = str(16-i)
                    table.cell(5,i)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "00FFFF")))
                table.cell(6,0).text = "Name"
                table.cell(7,0).text = "Type"
                table.cell(8,0).text = "Reset"

                for field_dict in self.excel_dict_new_sheet["register_dict"][register_name]["field_info_list"]:
                    #print(field_dict)
                    lsb = field_dict["lsb"]
                    msb = field_dict["lsb"] + field_dict["size"]-1
                    reset_value = field_dict['Reset Value']
                    for i in range(lsb,msb+1):
                        if i >= 16:
                            if field_dict["Customer_Visibility"] == "N":
                                table.cell(2,32-i).text = "Reserved"
                                table.cell(3,32-i).text = "RO"
                                table.cell(4,32-i).text = "0"
                                table.cell(2,32-i)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "c0c0c0")))
                                table.cell(3,32-i)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "c0c0c0")))
                                table.cell(4,32-i)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "c0c0c0")))
                            else:
                                table.cell(2,32-i).text = field_dict['Field Name'].upper()
                                table.cell(3,32-i).text = field_dict['Access Right'].replace("_E","").replace("_G","").replace("W1E","W1").replace("W0E","W0").replace("RC_RE","RC_R").upper()
                                table.cell(4,32-i).text = str(reset_value%2)
                        else:
                            if field_dict["Customer_Visibility"] == "N":
                                table.cell(6,16-i).text = "Reserved"
                                table.cell(7,16-i).text = "RO"
                                table.cell(8,16-i).text = "0"
                                table.cell(6,16-i)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "c0c0c0")))
                                table.cell(7,16-i)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "c0c0c0")))
                                table.cell(8,16-i)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "c0c0c0")))
                            else:
                                table.cell(6,16-i).text = field_dict['Field Name'].upper()
                                table.cell(7,16-i).text = field_dict['Access Right'].replace("_E","").replace("_G","").replace("W1E","W1").replace("W0E","W0").replace("RC_RE","RC_R").upper()
                                table.cell(8,16-i).text = str(reset_value%2)
                        reset_value = reset_value >> 1

                for i in range(1,17): #31~16
                    if table.cell(2,i).text == "":
                        table.cell(2,i).text = "Reserved"
                        table.cell(3,i).text = "RO"
                        table.cell(4,i).text = "0"
                        table.cell(2,i)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "c0c0c0")))
                        table.cell(3,i)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "c0c0c0")))
                        table.cell(4,i)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "c0c0c0")))
                for i in range(1,17): #15~0
                    if table.cell(6,i).text == "":
                        table.cell(6,i).text = "Reserved"
                        table.cell(7,i).text = "RO"
                        table.cell(8,i).text = "0"
                        table.cell(6,i)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "c0c0c0")))
                        table.cell(7,i)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "c0c0c0")))
                        table.cell(8,i)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'),color_value = "c0c0c0")))

                #merge the field  into one cell
                last_field_name = ""
                last_col_idx = 1
                for i in range(1,17): 
                    if table.cell(2,i).text == last_field_name:
                        table.cell(2,i).text = ""
                        table.cell(3,i).text = ""
                        table.cell(2,last_col_idx).merge(table.cell(2,i))
                        table.cell(3,last_col_idx).merge(table.cell(3,i))
                        table.cell(2,last_col_idx).text = table.cell(2,last_col_idx).text.strip()
                        table.cell(3,last_col_idx).text = table.cell(3,last_col_idx).text.strip()
                    else:
                        last_col_idx = i
                        last_field_name = table.cell(2,i).text
                last_field_name = ""
                last_col_idx = 1
                for i in range(1,17): 
                    if table.cell(6,i).text == last_field_name:
                        table.cell(6,i).text = ""
                        table.cell(7,i).text = ""
                        table.cell(6,last_col_idx).merge(table.cell(6,i))
                        table.cell(7,last_col_idx).merge(table.cell(7,i))
                        table.cell(6,last_col_idx).text = table.cell(6,last_col_idx).text.strip()
                        table.cell(7,last_col_idx).text = table.cell(7,last_col_idx).text.strip()
                    else:
                        table.cell(6,i).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        last_col_idx = i
                        last_field_name = table.cell(6,i).text

                #set alignment to center
                for i in range(1,17): 
                    table.cell(2,i).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table.cell(3,i).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table.cell(4,i).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table.cell(6,i).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table.cell(7,i).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    table.cell(8,i).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                document.add_paragraph('')
                document.add_paragraph(self.excel_dict_new_sheet["register_dict"][register_name]["Reg_Description"] , style='Normal')

                row_num = 1
                last_field_lsb = 32
                for field in self.excel_dict_new_sheet["register_dict"][register_name]["field_info_list"]:
                    if field["Customer_Visibility"] != "N": 
                        if(field["lsb"]+field["size"] != last_field_lsb): #this is reserved field
                            row_num = row_num + 1
                        row_num = row_num + 1
                        last_field_lsb = field["lsb"]  
                table = document.add_table(rows=row_num, cols=3,style = 'Table Grid')
                table.autofit = False
                table.allow_autofit = False
                widths =(Inches(1),Inches(1.5),Inches(4.5))
                for row in table.rows:
                    for idx,width in enumerate(widths):
                        row.cells[idx].width = width
                hdr_cells = table.rows[0].cells
                hdr_cells[0].paragraphs[0].add_run("Field").bold = True
                hdr_cells[1].paragraphs[0].add_run("Name").bold = True
                hdr_cells[2].paragraphs[0].add_run("Description").bold = True

                row_idx = 1
                last_field_lsb = 32
                for field in self.excel_dict_new_sheet["register_dict"][register_name]["field_info_list"]:
                    if field["Customer_Visibility"] != "N":
                        if(field["lsb"] + field["size"] != last_field_lsb ): #there is a reserved field
                            table.rows[row_idx].cells[0].text = "%d:%d"%(last_field_lsb -1 , field["lsb"] + field["size"])
                            table.rows[row_idx].cells[1].text = "Reserved"
                            table.rows[row_idx].cells[2].text = "Reserved field"
                            row_idx = row_idx + 1
                        table.rows[row_idx].cells[0].text = field["Field"].replace("[","").replace("]","")
                        table.rows[row_idx].cells[1].text = field["Field Name"].upper()
                        table.rows[row_idx].cells[2].text = field["Field_Description"]
                        last_field_lsb = field["lsb"] 
                        row_idx = row_idx + 1

                document.add_paragraph('')
                document.add_paragraph('')

            document.add_page_break()    
            document.save("%s/%s/%s_%s_reg_spec.docx"%(self.out_dir,sheet,module_name,sheet))
    #}}}

    def getTrimInfo(self):#{{{
        if not os.path.exists(self.opts.trim_cfg):
            self.log.fatal("file '%s' not existing, please check!!!" %(self.opts.trim_cfg))
            chipExitAll()

        fh = open(self.opts.trim_cfg, "r")
        trim_cfg = fh.readlines()
        fh.close()
        m = re.match("(\d+)'h(\w+)",trim_cfg[0].strip())
        if m:
            reset_width = m.group(1)
            reset_value_s = m.group(2)
        else:
            self.log.fatal("Wrong format in file '%s', reference: 345'habc12345678a please check!!!" %(self.opts.trim_cfg))
            chipExitAll()

        reset_value_l = []
        for s in reset_value_s:
            s = bin(int(s,16)).replace("0b","")
            for i in range(4 - len(s)):
                reset_value_l.append(0)
            for i in s:
                reset_value_l.append(int(i))
        reset_value_l.reverse()
        print(reset_value_l)
        return int(reset_width),reset_value_l

    #}}}

    def genRTL(self):#{{{
        self.log.debug("genRTL is running !")
        if self.opts.trim_cfg != "":
            reset_width,reset_value_l = self.getTrimInfo()
        else:
            reset_width = 0
            reset_value_l = []

        for sheet in sorted(self.excel_dict.keys()):
            chipMkDir("%s/%s"%(self.out_dir,sheet))
            module_name  = self.excel_dict[sheet]["title_dict"]["Module Name"]
            portlist = []
            if("Offset Format" in self.excel_dict[sheet]["title_dict"] and self.excel_dict[sheet]["title_dict"]["Offset Format"] == "Series"):
                addr_div = 1
            else:
                addr_div = self.excel_dict[sheet]["reg_width"]/8
            render_dict={"port_list":portlist,
                        "register_dict":self.excel_dict[sheet]["register_dict"],
                        "user": os.getlogin(),
                        "date":time.strftime("%Y-%m-%d %H:%M:%S",time.localtime()),
                        "address_msb":int(self.excel_dict[sheet]["title_dict"]["Address Width"].replace(".0",""))-1,
                        "module_name":self.excel_dict[sheet]["title_dict"]["Module Name"],
                        "sub_system":self.excel_dict[sheet]["title_dict"]["Sub-System"],
                        "reg_width":self.excel_dict[sheet]["reg_width"],
                        "addr_div":addr_div,
                        "reset_width":reset_width,
                        "reset_value_l":reset_value_l,
                        "version":self.excel_dict[sheet]["title_dict"]["Version"]
            }
            procTemplate("%s/%s/%s_rgf.v"%(self.out_dir,sheet,module_name),"reg_gen_rgf_template.v",render_dict)
    #}}}




    def mainProc(self): #{{{
        self.log.info("mainProc is running !")
        self.getExcelInfo()
        self.genRegModel()
        self.genRegModelLite()
        self.genRegDef()
        self.genRegDef2()
        self.genRegSvd()
        self.genRegSvdInternal()
        self.genSigBitInfo()
        self.genRTL()
        if self.doc:
            self.genDocx()
        pass
    #}}}
 #}}}

if __name__ == "__main__": #{{{

    args = myGetOpt()
    h_log = CCHIPLog(args.debug_en, vars(args), "reg_gen_%s"%(chipTimeStamp()))
    h_log.initPath("", "log", "reg_gen")
    dbg = args.debug_en

    proc = RegGen(args,h_log.log)
    proc.mainProc()

    h_log.log.info("%s Program Done by %s at %s !!!"%("reg_gen", chipUser(), chipDate()))

#}}}
# vim: fdm=marker

